#!/usr/bin/env python3
"""中诚咨询培训管理制度 - 基本配置版（精简版）"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21); section.page_height = Cm(29.7)
section.left_margin = Cm(2.8); section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5); section.bottom_margin = Cm(2)

C_DARK=RGBColor(0x1A,0x3A,0x5C); C_MID=RGBColor(0x2E,0x6D,0xA4)
C_LIGHT=RGBColor(0xD4,0xE6,0xF5); C_WHITE=RGBColor(0xFF,0xFF,0xFF)
C_TEXT=RGBColor(0x1F,0x1F,0x1F); C_GREY=RGBColor(0x60,0x60,0x60)

def sf(run,name='宋体',size=10.5,bold=False,color=None):
    run.font.name=name; run.font.size=Pt(size); run.font.bold=bold
    if color: run.font.color.rgb=color
def sp(p,before=2,after=5):
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
def shade(cell,hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex_color); tcPr.append(shd)
def h1(txt):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run(txt); sf(run,'黑体',16,bold=True,color=C_DARK)
    sp(p,before=16,after=6)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run2=p2.add_run('─'*50); sf(run2,'宋体',9,color=C_MID)
    sp(p2,before=0,after=4)
def h2(txt):
    p=doc.add_paragraph(); run=p.add_run('▌ '+txt)
    sf(run,'黑体',12,bold=True,color=C_DARK); sp(p,before=12,after=4)
def body(txt,indent=0):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(indent)
    run=p.add_run(txt); sf(run,'宋体',10.5,color=C_TEXT); sp(p,before=2,after=4)
def bul(txt,level=1):
    p=doc.add_paragraph()
    p.paragraph_format.left_indent=Cm(level*0.8)
    p.paragraph_format.first_line_indent=Cm(-0.4)
    run=p.add_run(f'●  {txt}'); sf(run,'宋体',10.5,color=C_TEXT); sp(p,before=2,after=3)
def tbl(headers,rows,widths=None,hdr='1A3A5C'):
    t=doc.add_table(rows=1+len(rows),cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; shade(c,hdr)
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: sf(r,'黑体',9.5,bold=True,color=C_WHITE)
            sp(p,before=3,after=3)
    for ri,row in enumerate(rows):
        bg='F0F5FA' if ri%2==0 else 'FFFFFF'
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; shade(c,bg)
            for p in c.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER if ci==0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs: sf(r,'宋体',9.5,color=C_TEXT)
                sp(p,before=3,after=3)
    if widths:
        for row in t.rows:
            for ci,w in enumerate(widths): row.cells[ci].width=Cm(w)
    return t
def pb(): doc.add_page_break()
def br(): doc.add_paragraph()
def div():
    p=doc.add_paragraph(); run=p.add_run('━'*55)
    sf(run,'宋体',8,color=C_LIGHT); sp(p,before=6,after=6)

# ===== 封面 =====
for _ in range(4): br()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=p.add_run('中诚智信工程咨询集团股份有限公司')
sf(run,'黑体',20,bold=True,color=C_DARK); sp(p,before=0,after=4)
p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
run2=p2.add_run('ZCEC GROUP'); run2.italic=True; sf(run2,'Arial',10,color=C_GREY); sp(p2,before=0,after=16)
div()
p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
run3=p3.add_run('培训管理制度')
sf(run3,'黑体',28,bold=True,color=C_DARK); sp(p3,before=12,after=6)
p4=doc.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
run4=p4.add_run('TRAINING MANAGEMENT POLICY  （Basic Edition）')
sf(run4,'Arial',12,bold=True,color=C_MID); sp(p4,before=0,after=20)
div()
br()
it=doc.add_table(rows=4,cols=4); it.style='Table Grid'; it.alignment=WD_TABLE_ALIGNMENT.CENTER
idata=[('文档编号','ZCEC-HR-2025-001-B','版本','V1.0（基本版）'),
       ('生效日期','2025年1月1日','下次评审','2026年1月'),
       ('制定部门','人力资源部','批准人','总经理'),
       ('适用范围','全集团','密级','内部公开')]
for ri,rd in enumerate(idata):
    for ci in range(4):
        c=it.rows[ri].cells[ci]; c.text=rd[ci]
        c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(c,'EAF2FB' if ci%2==0 else 'FFFFFF')
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if ci%2==0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                if ci%2==0: sf(r,'黑体',9,bold=True,color=C_DARK)
                else: sf(r,'宋体',9,color=C_TEXT)
            sp(p,before=3,after=3)
br(); p5=doc.add_paragraph(); p5.alignment=WD_ALIGN_PARAGRAPH.CENTER
run5=p5.add_run('— 保密提示：本文件为中诚咨询内部资料，请勿对外传播 —')
run5.italic=True; sf(run5,'宋体',9,color=C_GREY)
pb()

# ===== 目录 =====
h1('目  录')
for ch,title,page in [
    ('第一章','总则','1'),('第二章','培训体系','2'),('第三章','新员工培训','3'),
    ('第四章','在职培训','4'),('第五章','管理层培训','5'),('第六章','外部培训与证书','5'),
    ('第七章','效果评估','6'),('第八章','纪律与表单','6'),('第九章','附则','7')]:
    p=doc.add_paragraph()
    run=p.add_run(f'{ch}  {title}'); sf(run,'宋体',11,color=C_TEXT)
    run2=p.add_run(f'{" "*(46-len(ch)-len(title))} {page}'); sf(run2,'宋体',11,color=C_GREY)
    sp(p,before=3,after=3)
pb()

# ===== 第一章 总则 =====
h1('第一章  总则')
h2('第一条  目的')
body('规范公司培训管理工作，系统提升员工专业能力，支撑企业战略落地与规范化发展。')
h2('第二条  适用范围')
body('公司本部及全资、控股子公司全体员工。')
h2('第三条  培训方针')
bul('战略驱动：培训资源向公司战略重点领域倾斜')
bul('实战为本：培训内容紧密结合工作场景，学以致用')
bul('持续赋能：覆盖员工入、转、升、调全职业生命周期')
h2('第四条  培训目标（KPI）')
tbl(['KPI指标','目标值','责任部门'],
    [('年度培训覆盖率','100%','人力资源部'),
     ('人均年培训课时','≥40课时/人','人力资源部'),
     ('新员工岗位胜任周期','≤90天','人力资源部+用人部门'),
     ('培训满意度','≥4.2/5分','人力资源部'),
     ('关键岗位证书持有率','≥95%','人力资源部')],
    widths=[5,3,3])
pb()

# ===== 第二章 培训体系 =====
h1('第二章  培训体系')
h2('第五条  培训类别')
tbl(['类别','对象','周期','负责'],
    [('新员工培训','新入职员工','入职90天','人力资源部'),
     ('在职技能提升','全体在职员工','年度计划','人力资源部+业务部门'),
     ('管理层培训','部门负责人及以上','年度/半年度','人力资源部'),
     ('外部培训取证','相关岗位员工','按需申请','人力资源部统筹')],
    widths=[3,3.5,3,3.5])
h2('第六条  岗位序列')
tbl(['序列','代表岗位','核心培训内容'],
    [('专业技术','造价/监理/BIM工程师','清单计价、质量管控、BIM建模'),
     ('项目管理','全过程顾问、项目经理','项目管理、合同管理、设计管理'),
     ('招标代理','招标专员/经理','招投标法规、流程合规'),
     ('职能支持','财务/人事/行政/法务','专业技能、上市合规'),
     ('高层管理','总经理/副总/总监','战略规划、领导力')],
    widths=[2.5,4,6.5])
h2('第七条  培训方式')
bul('内部集中培训：系统性强，成本低，统一管理')
bul('岗位导师制：个性化辅导，一对一实战传承')
bul('在线学习（E-Learning）：灵活便捷，碎片时间')
bul('外派进修：高管培训、资质取证、标杆考察')
h2('第八条  年度培训日历')
tbl(['月份','固定培训','备注'],
    [('1月','年度计划发布会','各部门提报次年需求'),
     ('2月','开工安全质量培训','全员'),
     ('3月','新员工90天启动','BIM专项'),
     ('5月','全过程咨询训练营','项目序列'),
     ('6月','招投标合规月','安全生产月'),
     ('8月','中层管理培训','读书会'),
     ('9月','新员工结业答辩','内训师评审'),
     ('11月','培训满意度调研','次年需求调研'),
     ('12月','年度总结表彰','次年计划发布')],
    widths=[1.5,5.5,6])
pb()

# ===== 第三章 新员工培训 =====
h1('第三章  新员工培训（90天）')
h2('第九条  培养阶段')
tbl(['阶段','周期','核心任务','方式','考核'],
    [('通识培训','第1周','文化融入、制度知晓','集中授课','笔试≥60分'),
     ('轮岗认知','第2-4周','主要业务板块了解','各部门轮转','心得报告'),
     ('岗位实训','第5-12周','本岗位核心技能','导师制+项目实战','周报+月度考核'),
     ('定岗考核','第13周','综合评定','答辩','≥60分定岗')],
    widths=[2,1.8,3.5,2.8,3])
h2('第十条  导师制')
bul('入职第5周前为每位新员工指定导师（原则上同部门高级工程师/项目经理）')
bul('导师职责：制定IDP、每周辅导≥1次、每月填写《导师辅导记录表》')
bul('导师补贴：200元/人/月；年度优秀导师评选，颁发荣誉证书+奖金500元')
h2('第十一条  定岗考核')
bul('第85天前提交《90天培训总结报告》（≥3000字）')
bul('答辩委员会：人力资源部+直线经理+导师')
bul('评分：业务理解30% + 技能掌握40% + 态度素养20% + 表达能力10%')
bul('考核不合格者给予一次补考机会，仍不合格视为不符合录用条件')
h2('第十二条  淘汰机制')
body('试用期出现以下情形，公司有权依法解除劳动合同：')
tbl(['情形','处理'],
    [('考核不合格（补考后仍＜60分）','解除，无经济补偿'),
     ('严重违反公司规章制度','解除，无经济补偿'),
     ('入职材料虚假陈述','解除，无经济补偿'),
     ('旷工连续3天或累计5天以上','解除，支付已工作期间工资')],
    widths=[5.5,7.5])
pb()

# ===== 第四章 在职培训 =====
h1('第四章  在职技能提升培训')
h2('第十三条  年度培训计划制定')
body('每年12月各部门提报次年培训需求，人力资源部汇总分析后于次年1月15日前发布年度培训计划。')
h2('第十四条  重点培训项目')
tbl(['培训项目','对象','课时','频次','目标'],
    [('清单计价与定额应用','造价工程师','16课时','每年2次','持证率100%'),
     ('全过程工程咨询','项目序列','24课时','每年1次','核心人员轮训'),
     ('BIM建模专项','技术序列','32课时','每年1次','持BIM证书率80%'),
     ('招投标合规','招标序列+全员','8课时','每年1次','合规零违规'),
     ('安全生产质量','监理+现场','12课时','每季度','安全事故为零'),
     ('职业素养','全员','4课时','每年1次','客户满意度提升')],
    widths=[3.5,3,1.5,2,3])
h2('第十五条  内训师制度')
tbl(['级别','申报条件','课酬','年度任务'],
    [('初级讲师','主管级+自愿申请','200元/课时','≥20课时'),
     ('中级讲师','初级满1年+60课时+评价≥4.0','400元/课时','≥30课时'),
     ('高级讲师','中级满1年+120课时+体系化课程','600元/课时','≥40课时')],
    widths=[2.5,6,2.5,3])
h2('第十六条  培训积分')
body('公司建立培训积分制度，年度A类必修积分要求：普通员工≥40分，管理层≥24分，未达标者年度绩效不得评为"优秀"。')
h2('第十七条  案例萃取')
body('各部门每季度萃取≥1个典型案例，提交至公司案例库。优秀案例采纳奖励500元/个，用于内训课程和新员工培训。')
pb()

# ===== 第五章 管理层培训 =====
h1('第五章  管理层培训')
h2('第十八条  培训对象与内容')
tbl(['层级','对象','培训重点','课时/年'],
    [('决策层','总经理/副总/总监','战略规划、资本运营、行业政策','≥48课时'),
     ('中层','部门负责人','目标管理、绩效辅导、跨部门协作','≥36课时'),
     ('后备','高潜员工','职业素养、角色转换、基础管理','≥24课时')],
    widths=[1.8,4,6.2,2])
h2('第十九条  培训方式')
bul('年度集中培训：每年1-2次管理层集中培训（2-3天/次）')
bul('外派高端研修：选派优秀管理人员参加EMBA、总裁班、行业论坛')
bul('标杆企业考察：每年组织1-2次行业标杆企业参访')
bul('管理层读书会：每月至少读1本管理类书籍，每季度举办读书分享会')
h2('第二十条  外训审批')
tbl(['费用范围','审批人','服务期要求'],
    [('<5000元','人力资源部负责人','—'),
     ('5000-10000元','人力资源部+分管领导','—'),
     ('≥10000元','总经理办公会','≥3年'),
     ('EMBA/总裁班','总经理特批','5年（退50%）')],
    widths=[3,5,5])
pb()

# ===== 第六章 外部培训与证书 =====
h1('第六章  外部培训与证书管理')
h2('第二十一条  证书分类与补贴')
tbl(['证书类型','报销上限','一次性奖励','月度补贴','服务期'],
    [('一级造价工程师','8000元','5000元','500元/月','3年'),
     ('一级建造师','8000元','5000元','500元/月','3年'),
     ('监理工程师','6000元','5000元','500元/月','3年'),
     ('招标师','4000元','3000元','300元/月','2年'),
     ('BIM高级建模师','3000元','2000元','200元/月','2年')],
    widths=[3.5,2.5,2.5,2.5,3])
h2('第二十二条  服务期违约')
tbl(['情形','处理'],
    [('服务期满','服务期届满，证书权益保留'),
     ('期内主动离职','按比例退还：应退金额=报销总额×（未服务月数/规定月数）'),
     ('公司依法解除（员工无过错）','无需退还'),
     ('严重违规被解除','须全额退还，不享受任何证书补贴')],
    widths=[5,8])
h2('第二十三条  继续教育')
body('执业资格证书的继续教育费用由公司全额报销，人力资源部建立证书有效期台账，提前6个月提醒员工参加继续教育。')
pb()

# ===== 第七章 效果评估 =====
h1('第七章  培训效果评估')
h2('第二十四条  柯氏四级评估')
tbl(['评估层级','评估内容','评估方式','评估时点'],
    [('第一级：反应层','学员满意度','问卷（5分制）','培训结束时'),
     ('第二级：学习层','知识技能掌握','笔试+实操考核','培训结束+30天'),
     ('第三级：行为层','训后行为改变','直线经理评价','培训后90天'),
     ('第四级：结果层','对业绩的贡献','绩效数据对比','培训后6-12个月')],
    widths=[3,3,4,3])
h2('第二十五条  评估结果应用')
bul('评估结果记入员工培训档案，作为晋升、评优的重要参考')
bul('培训满意度连续2期低于3.5分的课程，暂停开设并优化改进')
bul('培训后行为转化率达标的员工，优先安排下一阶段培训')
h2('第二十六条  培训满意度调查')
body('每期培训结束时发放满意度问卷，评估维度：课程实用性（25%）、讲师质量（20%）、组织管理（15%）、学习收获（20%）、建议意见（20%）。')
pb()

# ===== 第八章 纪律与表单 =====
h1('第八章  培训纪律与表单')
h2('第二十七条  出勤纪律')
tbl(['情形','处理'],
    [('迟到≤10分钟','口头警告，计入出勤'),
     ('迟到>10分钟','扣罚当月绩效50元，不记学时'),
     ('旷课（无故）','扣罚当月绩效200元/课时，书面警告'),
     ('伪造培训记录','解除劳动合同，追究法律责任')],
    widths=[5,8])
h2('第二十八条  请假审批')
tbl(['培训时长','审批人','补训要求'],
    [('≤4课时','直线经理','需补训'),
     ('4-16课时','直线经理+人力资源部','需补训'),
     ('>16课时（外训）','直线经理+人力资源部+分管领导','需补训')],
    widths=[3.5,5.5,4])
h2('第二十九条  配套表单清单（18张）')
tbl(['编号','表单名称','用途'],
    [('HR-001','新员工入职培训记录表','记录新员工通识培训全过程'),
     ('HR-002','员工年度培训需求申请表','各部门提报次年培训需求'),
     ('HR-003','外部培训申请表','员工申请外训'),
     ('HR-004','培训请假单','员工因故无法参加培训'),
     ('HR-005','培训满意度调查问卷','每期培训结束收集'),
     ('HR-006','培训考核评分表','记录笔试/实操成绩'),
     ('HR-007','员工培训档案表','一人一档，动态更新'),
     ('HR-008','内部讲师课酬申请表','内训师申请课酬'),
     ('HR-009','证书补贴申请表','员工取证后申领补贴'),
     ('HR-010','导师辅导记录表','导师记录辅导期沟通'),
     ('HR-011','培训效果评估表（柯氏四级）','培训后30/90天效果追踪'),
     ('HR-012','新员工周报','新员工每周汇报学习进展'),
     ('HR-013','新员工IDP计划表','个人发展计划'),
     ('HR-014','新员工培训总结报告','90天培训成果汇报'),
     ('HR-015','年度培训计划表','各部门提报年度培训计划'),
     ('HR-016','培训工作月报','每月培训工作总结'),
     ('HR-017','培训签到表','每次培训记录出席情况'),
     ('HR-018','培训费用报销单','培训费用报销申请')],
    widths=[2,4,7])
pb()

# ===== 第九章 附则 =====
h1('第九章  附则')
h2('第三十条  解释权')
body('本制度由人力资源部负责解释。执行过程中如有争议，报总经理办公会裁决。')
h2('第三十一条  修订机制')
bul('年度修订：每年12月启动，次年1月发布')
bul('即时修订：法规变化、战略调整、重大漏洞时可即时修订')
h2('第三十二条  制度配套')
body('本制度与以下制度配套使用：《员工手册》《绩效管理制度》《薪酬管理制度》《晋升管理制度》《财务报销管理制度》。')
h2('第三十三条  生效')
body('本制度自2025年1月1日起施行，原V0.X版同步废止。')
br(); div(); br()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
run=p.add_run('中诚智信工程咨询集团股份有限公司  人力资源部  2025年1月')
sf(run,'宋体',10,color=C_GREY); sp(p,before=4,after=4)
p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
run2=p2.add_run('— 本制度配套表单共18张，编号HR-001至HR-018 —')
run2.italic=True; sf(run2,'宋体',9,color=C_GREY)

output='/Users/huahaha/Desktop/中诚咨询培训管理制度/中诚咨询培训管理制度_基本配置版.docx'
doc.save(output)
print(f'✅ 基本版已生成: {output}')
