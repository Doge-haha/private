#!/usr/bin/env python3
"""生成中诚咨询全套培训管理制度 Word文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ========== 页面设置 ==========
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2)

# ========== 辅助函数 ==========
def set_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, '黑体', 16, bold=True)
    p.space_before = Pt(24)
    p.space_after = Pt(12)
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '黑体', 14, bold=True)
    p.space_before = Pt(18)
    p.space_after = Pt(8)
    return p

def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '黑体', 12, bold=True)
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, '宋体', 11)
    p.space_before = Pt(3)
    p.space_after = Pt(3)
    return p

def bullet(text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.8)
    run = p.add_run(text)
    set_font(run, '宋体', 11)
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, '黑体', 10, bold=True)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, '宋体', 10)
    return table

# ========== 封面 ==========
doc.add_paragraph()
doc.add_paragraph()
heading1("中诚智信工程咨询集团股份有限公司")
doc.add_paragraph()
heading1("培 训 管 理 制 度")
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("（2025版）")
set_font(run, '宋体', 14)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = ['文件编号', 'ZCEC-HR-2025-001', '版本', 'V1.0', '生效日期', '2025年1月1日', '适用范围', '全集团（含子公司）', '制定部门', '人力资源部']
for i in range(5):
    info_table.rows[i].cells[0].text = labels[i*2]
    info_table.rows[i].cells[1].text = labels[i*2+1]
    for cell in info_table.rows[i].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, '宋体', 11)

doc.add_page_break()

# ========== 目录 ==========
heading1("目  录")
toc_items = [
    ("第一章", "总则"),
    ("第二章", "培训组织与职责"),
    ("第三章", "培训分类与体系"),
    ("第四章", "新员工培训"),
    ("第五章", "在职技能提升培训"),
    ("第六章", "管理层培训"),
    ("第七章", "外部培训与证书管理"),
    ("第八章", "培训效果评估"),
    ("第九章", "培训纪律与考核"),
    ("第十章", "培训档案管理"),
    ("第十一章", "附则"),
    ("附件", "培训表单清单"),
]
for ch, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{ch}  {title}")
    set_font(run, '宋体', 11)
    p.space_before = Pt(4)
    p.space_after = Pt(4)

doc.add_page_break()

# ========== 正文 ==========
heading1("第一章  总则")

heading2("第一条  目的")
body('为建立健全中诚智信工程咨询集团股份有限公司（以下简称"公司"）培训管理体系，规范员工培训行为，提升员工专业技能和综合素质，增强企业核心竞争力，结合公司实际情况，特制定本制度。')

heading2("第二条  适用范围")
body("本制度适用于公司本部及全资子公司、控股子公司全体员工。")
body("参股公司参照本制度执行，或另行制定相关规定。")

heading2("第三条  制定依据")
body("本制度依据以下法律法规和规范性文件制定：")
bullet("《中华人民共和国劳动合同法》")
bullet("《中华人民共和国职业教育法》")
bullet("《企业职工培训规定》（劳部发〔1996〕370号）")
bullet("《关于加强企业职工教育培训工作的意见》（国资厅发〔2010〕45号）")
bullet("《中诚智信工程咨询集团股份有限公司章程》")

heading2("第四条  培训原则")
bullet("战略导向原则：培训服务于公司发展战略和业务需要")
bullet("全员覆盖原则：培训覆盖全体员工，按岗位需求差异化实施")
bullet("实效性原则：注重培训效果，强调学以致用")
bullet("持续性原则：建立长期、系统的培训机制")

heading2("第五条  培训目标")
bullet("员工专业技能达标率：95%以上")
bullet("新员工岗位胜任周期：不超过90天")
bullet("关键岗位证书持有率：100%")
bullet("年度培训覆盖率：100%")

# ========== 第二章 ==========
heading1("第二章  培训组织与职责")

heading2("第六条  人力资源部职责")
bullet("负责制定公司年度培训计划并组织实施")
bullet("建立和维护培训管理体系，制定相关制度")
bullet("组织开发内部培训课程和教材")
bullet("管理培训档案和培训费用")
bullet("对子公司培训工作进行指导、监督和考核")

heading2("第七条  各部门职责")
bullet("配合人力资源部开展本部门培训工作")
bullet("推荐内部培训讲师，提供业务技术支持")
bullet("负责本部门员工在岗培训的跟踪和评估")
bullet("按时提交本部门年度培训需求和培训记录")

heading2("第八条  直线经理职责")
bullet("了解本部门员工的培训需求")
bullet("对下属参加培训进行审批并跟踪培训效果")
bullet("配合人力资源部对下属进行岗位技能考核")

heading2("第九条  员工权利与义务")
body("员工有权利参加公司组织的各类培训活动，享有培训费用报销的权利。")
body("员工有义务遵守培训纪律，完成培训任务，将所学知识应用于工作实践。")

# ========== 第三章 ==========
heading1("第三章  培训分类与体系")

heading2("第十条  培训体系架构")
body("公司培训体系分为四大类别：新员工培训、在职技能提升培训、管理层培训、外部培训与取证。")

add_table(
    ['培训类别', '培训对象', '负责部门', '开展周期'],
    [
        ['新员工培训', '新入职员工', '人力资源部', '入职即启动'],
        ['在职技能提升', '全体在职员工', '人力资源部+业务部门', '季度/月度'],
        ['管理层培训', '部门负责人及以上', '人力资源部', '年度/半年度'],
        ['外部培训取证', '相关岗位员工', '人力资源部', '按需'],
    ]
)
doc.add_paragraph()

heading2("第十一条  岗位序列划分")
add_table(
    ['序列', '岗位方向', '核心培训内容'],
    [
        ['专业技术序列', '造价工程师、监理工程师、BIM工程师', '清单计价、定额套用、质量管控、BIM建模'],
        ['项目管理序列', '全过程咨询顾问、项目经理', '项目管理、合同管理、设计管理'],
        ['招标代理序列', '招标专员、招标项目经理', '招投标法规、流程合规、文件编制'],
        ['职能管理序列', '财务、人事、行政、法务', '专业技能、上市合规、职业素养'],
        ['高层管理序列', '总经理、副总经理、总监', '战略规划、行业政策、领导力'],
    ]
)
doc.add_paragraph()

# ========== 第四章 ==========
heading1("第四章  新员工培训")

heading2("第十二条  培训目标")
bullet("帮助新员工了解公司文化、战略和组织架构")
bullet("掌握岗位必备的基础知识和基本技能")
bullet("快速融入团队，缩短岗位胜任周期")

heading2("第十三条  培训周期")
body("新员工培训周期为90天，分为四个阶段：")

add_table(
    ['阶段', '时间', '培训内容', '考核方式'],
    [
        ['第一阶段：通识培训', '入职第1周', '企业文化、规章制度、安全教育、行业基础知识', '笔试'],
        ['第二阶段：轮岗认知', '入职第2-4周', '造价/招标/监理/全过程业务板块轮转了解', '心得报告'],
        ['第三阶段：岗位实训', '入职第2-3月', '导师制，岗位基础技能实操', '实操考核'],
        ['第四阶段：定岗考核', '入职第4月', '综合考核评定', '答辩/考核定岗'],
    ]
)
doc.add_paragraph()

heading2("第十四条  导师制")
body("新员工入职后，公司为其指定一名导师（一般为同部门业务骨干），导师负责：")
bullet("业务指导和技能传授")
bullet("工作习惯和职业素养引导")
bullet("定期沟通，了解员工思想动态")
body("导师辅导期不少于3个月，辅导期间享受导师补贴。")

heading2("第十五条  培训纪律")
bullet("新员工必须参加入职培训的全部课程，因故缺勤需提前请假并补训")
bullet("培训考核不合格者，给予一次补考机会；补考仍不合格者，视为不符合录用条件")

# ========== 第五章 ==========
heading1("第五章  在职技能提升培训")

heading2("第十六条  培训形式")
bullet("内部培训：公司内部讲师授课、案例分析、经验分享")
bullet("岗位轮训：跨部门或跨业务条线的实践学习")
bullet("在线学习：依托网络平台自主学习")
bullet("项目实践：以实际项目为载体的在岗学习")

heading2("第十七条  年度培训计划")
body("每年12月，人力资源部向各部门征集次年培训需求，编制《年度培训计划》，经总经理审批后于次年1月公布实施。")

heading2("第十八条  内部讲师制度")
body("公司建立内部讲师库，各部门推荐业务骨干担任内部讲师。内部讲师享有课酬补贴：")
add_table(
    ['讲师级别', '课酬标准（元/课时）', '晋升条件'],
    [
        ['初级讲师', '200', '累计授课30课时以上，经评审晋升'],
        ['中级讲师', '400', '累计授课80课时以上，学员评价良好'],
        ['高级讲师', '600', '累计授课150课时以上，形成体系化课程'],
    ]
)
doc.add_paragraph()

heading2("第十九条  专项培训")
body("根据业务需要，公司开展以下专项培训：")
bullet("BIM技术应用培训：面向造价和设计人员，每年至少1次")
bullet("全过程工程咨询培训：面向项目管理人员，每年至少2次")
bullet("招投标法规更新培训：政策变化时及时组织，全员覆盖")
bullet("安全生产与质量管理培训：面向监理和现场管理人员，每季度1次")

# ========== 第六章 ==========
heading1("第六章  管理层培训")

heading2("第二十条  培训对象")
body("培训对象包括：部门负责人、项目总监、副总经理、总经理及后备管理人才。")

heading2("第二十一条  培训内容")
bullet("宏观经济与行业发展趋势分析")
bullet("企业战略规划与执行")
bullet("上市合规与公司治理")
bullet("领导力与管理技能提升")
bullet("行业最新政策法规解读")

heading2("第二十二条  培训方式")
bullet("年度集中培训：每年至少组织1次管理层集中培训")
bullet("外派学习：选派优秀管理人员参加外部高端培训")
bullet("标杆企业考察：定期组织赴行业标杆企业交流学习")
bullet("读书分享会：管理层每月至少读一本管理类书籍并分享")

# ========== 第七章 ==========
heading1("第七章  外部培训与证书管理")

heading2("第二十三条  外部培训申请")
body("员工参加外部培训，需填写《外部培训申请表》，经直线经理、人力资源部及分管领导审批后方可参加。")

heading2("第二十四条  培训费用报销")
add_table(
    ['证书类型', '公司承担上限', '服务期要求', '违约处理'],
    [
        ['一级建造师', '8000元/人', '取得后服务满3年', '未满服务期离职，退还全部培训费'],
        ['一级造价工程师', '6000元/人', '取得后服务满3年', '同上'],
        ['监理工程师', '4000元/人', '取得后服务满2年', '同上'],
        ['其他执业资格', '3000元/人', '取得后服务满2年', '同上'],
    ]
)
doc.add_paragraph()

heading2("第二十五条  证书补贴")
body("员工取得公司所需执业资格证书后，享受一次性奖励补贴，并在月度薪资中发放技能补贴：")
bullet("一级建造师：一次性奖励5000元，月度技能补贴500元")
bullet("一级造价工程师：一次性奖励4000元，月度技能补贴400元")
bullet("监理工程师：一次性奖励3000元，月度技能补贴300元")

heading2("第二十六条  继续教育")
body("执业资格证书的继续教育费用由公司全额报销。员工应合理安排继续教育学习，确保证书有效性。")

# ========== 第八章 ==========
heading1("第八章  培训效果评估")

heading2("第二十七条  评估方式")
body("公司采用柯氏四级评估模型（Kirkpatrick Model）进行培训效果评估：")

add_table(
    ['评估层级', '评估内容', '评估方式', '评估时点'],
    [
        ['第一级：反应层', '学员满意度', '满意度调查问卷', '培训结束时'],
        ['第二级：学习层', '知识技能掌握程度', '笔试、实操考核', '培训结束时/培训后30天'],
        ['第三级：行为层', '训后行为改变', '直线经理评价、360度评估', '培训后90天'],
        ['第四级：结果层', '对业绩的影响', '绩效数据对比', '培训后6-12个月'],
    ]
)
doc.add_paragraph()

heading2("第二十八条  培训考核结果应用")
bullet("考核结果记入员工培训档案")
bullet("考核优秀者在晋升、评优中优先考虑")
bullet("考核不合格者须参加补训或补考")

# ========== 第九章 ==========
heading1("第九章  培训纪律与考核")

heading2("第二十九条  培训出勤纪律")
bullet("培训期间不得迟到、早退、旷课，因故不能参加须提前办理请假手续")
bullet("培训期间应认真听讲，积极参与互动，禁止做与培训无关的事情")
bullet("违反培训纪律者，视情节轻重给予批评教育或扣减绩效")

heading2("第三十条  培训请假")
body("员工参加培训需请假的，应提前填写《培训请假单》，按照以下权限审批：")
bullet("培训时间1天以内：直线经理审批")
bullet("培训时间1天以上：直线经理+人力资源部审批")

heading2("第三十一条  培训档案考核")
body("人力资源部对每位员工建立《员工培训档案》，记录内容包括：")
bullet("培训记录：课程名称、培训时间、时长、考核成绩")
bullet("证书记录：取得证书名称、取得时间、有效期")
bullet("培训费用：本人发生的培训费用总额")

# ========== 第十章 ==========
heading1("第十章  培训档案管理")

heading2("第三十二条  档案内容")
body("员工培训档案包括：")
bullet("新员工培训记录及考核表")
bullet("在职培训记录及考核成绩")
bullet("外部培训申请表及费用报销凭证")
bullet("执业资格证书复印件")
bullet("培训满意度调查问卷")

heading2("第三十三条  档案管理")
bullet("培训档案由人力资源部统一管理，保存期限不少于3年")
bullet("员工离职时，培训档案随人事档案一并转移或按规定年限保存")
bullet("电子化档案与纸质档案并行管理，确保信息安全")

# ========== 第十一章 ==========
heading1("第十一章  附则")

heading2("第三十四条  解释权")
body("本制度由人力资源部负责解释。")

heading2("第三十五条  修订")
body("本制度每年修订一次，遇下列情况可即时修订：")
bullet("公司组织架构重大调整")
bullet("国家或地方相关法律法规修订")
bullet("公司战略方向重大变化")

heading2("第三十六条  生效")
body("本制度自2025年1月1日起施行。")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("中诚智信工程咨询集团股份有限公司")
set_font(run, '宋体', 11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("人力资源部")
set_font(run, '宋体', 11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("2025年1月")
set_font(run, '宋体', 11)

doc.add_page_break()

# ========== 附件 ==========
heading1("附件：培训表单清单")

forms = [
    ("HR-001", "新员工入职培训记录表", "新员工完成入职培训后填写"),
    ("HR-002", "员工培训需求申请表", "员工申请培训时填写"),
    ("HR-003", "外部培训申请表", "员工申请外训时填写"),
    ("HR-004", "培训请假单", "员工因故无法参加培训时填写"),
    ("HR-005", "培训满意度调查问卷", "每次培训结束后收集"),
    ("HR-006", "培训考核评分表", "培训结束时进行考核评分"),
    ("HR-007", "员工培训档案表", "人力资源部建立，一人一档"),
    ("HR-008", "内部讲师课酬申请表", "内部讲师申请课酬时填写"),
    ("HR-009", "证书补贴申请表", "员工取得证书后申请补贴时填写"),
    ("HR-010", "年度培训计划表", "每年12月各部门填写提交"),
    ("HR-011", "培训效果评估表（柯氏四级）", "培训后30/90天进行效果追踪"),
    ("HR-012", "导师辅导记录表", "导师记录辅导期内与新员工的沟通情况"),
]

add_table(
    ['表单编号', '表单名称', '用途说明'],
    forms
)

output_path = '/Users/huahaha/Downloads/中诚咨询培训管理制度_2025.docx'
doc.save(output_path)
print(f"✅ 文档已生成: {output_path}")
