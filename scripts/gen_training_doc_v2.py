#!/usr/bin/env python3
"""生成中诚咨询全套培训管理制度 Word文档 - 专业版"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ========== 页面设置 ==========
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2)

# ========== 颜色定义 ==========
BRAND_DARK = RGBColor(0x1A, 0x3A, 0x5C)    # 深蓝
BRAND_MID  = RGBColor(0x2E, 0x6D, 0xA4)    # 中蓝
BRAND_LIGHT= RGBColor(0xD4, 0xE6, 0xF5)    # 浅蓝背景
ACCENT     = RGBColor(0xE8, 0x6A, 0x1A)    # 橙色强调
TEXT_DARK  = RGBColor(0x1F, 0x1F, 0x1F)
TEXT_GREY  = RGBColor(0x60, 0x60, 0x60)

# ========== 辅助函数 ==========
def set_font(run, name='宋体', size=12, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def set_para_spacing(p, before=0, after=6, line=1.5):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), kwargs[edge].get('val','single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[edge].get('color','auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def heading1(text, add_line=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, '黑体', 18, bold=True, color=BRAND_DARK)
    set_para_spacing(p, before=20, after=8)
    if add_line:
        # 分隔线
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run('─' * 42)
        set_font(run2, '宋体', 9, color=BRAND_MID)
        set_para_spacing(p2, before=0, after=4)
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run('▌ ' + text)
    set_font(run, '黑体', 13, bold=True, color=BRAND_DARK)
    set_para_spacing(p, before=16, after=6)
    # 左侧色条
    p.paragraph_format.left_indent = Cm(0)
    return p

def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run('▶ ' + text)
    set_font(run, '黑体', 11, bold=True, color=BRAND_MID)
    set_para_spacing(p, before=10, after=4)
    return p

def body(text, indent=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    if bold_prefix:
        run1 = p.add_run(bold_prefix)
        set_font(run1, '宋体', 10.5, bold=True, color=TEXT_DARK)
        run2 = p.add_run(text)
        set_font(run2, '宋体', 10.5, color=TEXT_DARK)
    else:
        run = p.add_run(text)
        set_font(run, '宋体', 10.5, color=TEXT_DARK)
    set_para_spacing(p, before=2, after=4)
    return p

def bullet(text, level=1, marker='●'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(level * 0.8)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f'{marker}  {text}')
    set_font(run, '宋体', 10.5, color=TEXT_DARK)
    set_para_spacing(p, before=2, after=3)
    return p

def numbered(text, level=1, num=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(level * 0.8)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f'{num or "①②③④⑤⑥⑦⑧⑨⑩"[level-1] if level<=10 else str(level)+"."}  {text}')
    set_font(run, '宋体', 10.5, color=TEXT_DARK)
    set_para_spacing(p, before=2, after=3)
    return p

def divider():
    p = doc.add_paragraph()
    run = p.add_run('━' * 50)
    set_font(run, '宋体', 8, color=BRAND_LIGHT)
    set_para_spacing(p, before=6, after=6)

def add_table(headers, rows, col_widths=None, header_color='1A3A5C'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, header_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, '黑体', 10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            set_para_spacing(p, before=4, after=4)
    # Data
    for ri, row in enumerate(rows):
        bg = 'F0F5FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade_cell(cell, bg)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_font(run, '宋体', 10, color=TEXT_DARK)
                set_para_spacing(p, before=3, after=3)
    # Col widths
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return table

def info_box(title, items, bg='EAF2FB', title_color='1A3A5C'):
    """信息框"""
    p = doc.add_paragraph()
    run = p.add_run('▎' + title)
    set_font(run, '黑体', 11, bold=True, color=RGBColor.from_string(title_color))
    set_para_spacing(p, before=10, after=2)
    for item in items:
        bullet(item)

def page_break():
    doc.add_page_break()

# ========== 封面 ==========
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('中诚智信工程咨询集团股份有限公司')
set_font(run, '黑体', 22, bold=True, color=BRAND_DARK)
set_para_spacing(p, before=0, after=6)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('ZC Intel Tec Engineering Consulting Group Co.,Ltd.')
set_font(run2, 'Arial', 10, italic=True, color=TEXT_GREY)
set_para_spacing(p2, before=0, after=20)

divider()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('培 训 管 理 制 度')
set_font(run3, '黑体', 30, bold=True, color=BRAND_DARK)
set_para_spacing(p3, before=16, after=8)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('TRAINING MANAGEMENT POLICY')
set_font(run4, 'Arial', 13, bold=True, color=BRAND_MID)
set_para_spacing(p4, before=0, after=30)

divider()

for _ in range(3):
    doc.add_paragraph()

# 文档信息表
info_table = doc.add_table(rows=6, cols=4)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('文档编号', 'ZCEC-HR-2025-001', '版本', 'V1.0'),
    ('生效日期', '2025年1月1日', '下次评审', '2026年1月'),
    ('制定部门', '人力资源部', '批准人', '总经理'),
    ('适用范围', '全集团（含子公司）', '页数', '约32页'),
    ('密级', '内部公开', '制度级别', '公司级'),
    ('制定日期', '2025年1月', '修订日期', '—'),
]
for ri, row_data in enumerate(info_data):
    for ci in range(4):
        cell = info_table.rows[ri].cells[ci]
        cell.text = row_data[ci]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, 'EAF2FB' if ci % 2 == 0 else 'FFFFFF')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                if ci % 2 == 0:
                    set_font(run, '黑体', 9, bold=True, color=BRAND_DARK)
                else:
                    set_font(run, '宋体', 9, color=TEXT_DARK)
            set_para_spacing(p, before=3, after=3)

for _ in range(2):
    doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run('— 保密提示：本文件为中诚咨询内部资料，请勿对外传播 —')
set_font(run5, '宋体', 9, italic=True, color=TEXT_GREY)

page_break()

# ========== 目录 ==========
heading1('目  录')

toc_items = [
    ('第一章', '总则', '1'),
    ('第二章', '培训组织与职责', '3'),
    ('第三章', '培训体系架构', '4'),
    ('第四章', '新员工培训（90天培养计划）', '6'),
    ('第五章', '在职技能提升培训', '9'),
    ('第六章', '管理层培训', '11'),
    ('第七章', '外部培训与证书管理', '13'),
    ('第八章', '培训效果评估体系', '15'),
    ('第九章', '培训纪律与考核', '17'),
    ('第十章', '培训档案管理', '18'),
    ('第十一章', '年度培训实施计划表', '19'),
    ('第十二章', '表单清单与填写说明', '22'),
    ('第十三章', '附则', '25'),
]

for ch, title, page in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f'{ch}  {title}')
    set_font(run, '宋体', 11, color=TEXT_DARK)
    # 页码点线
    dots = ' ' * max(1, 48 - len(ch) - len(title))
    run2 = p.add_run(f'{dots} {page}')
    set_font(run2, '宋体', 11, color=TEXT_GREY)
    set_para_spacing(p, before=3, after=3)

page_break()

# ========== 第一章 ==========
heading1('第一章  总则')

heading2('第一条  目的与意义')
body('为建立健全中诚智信工程咨询集团股份有限公司（以下简称"公司"）培训管理体系，系统提升员工专业能力与职业素养，支撑公司战略落地与上市后的规范化发展，结合公司实际情况，特制定本制度。')

heading2('第二条  适用范围')
body('本制度适用于公司本部及全资子公司、控股子公司全体员工。参股公司参照执行。')

heading2('第三条  制定依据')
bullet('《中华人民共和国劳动合同法》')
bullet('《中华人民共和国职业教育法》')
bullet('《企业职工培训规定》（劳部发〔1996〕370号）')
bullet('《中诚智信工程咨询集团股份有限公司章程》')
bullet('《北京证券交易所股票上市规则》')

heading2('第四条  培训方针')
body('公司培训坚持"战略驱动、实战为本、持续赋能"的方针：', bold_prefix='战略驱动：')
bullet('培训服务于公司发展战略和业务需要')
bullet('重点支撑全过程工程咨询、BIM技术等核心能力的构建')
body('', bold_prefix='实战为本：')
bullet('培训内容紧密结合工作场景，学以致用')
bullet('以真实项目案例为载体，提升解决实际问题的能力')
body('', bold_prefix='持续赋能：')
bullet('建立长期、系统的培训机制，覆盖员工全职业生命周期')
bullet('鼓励员工自主学习，形成终身学习的组织文化')

heading2('第五条  培训目标（KPI）')
add_table(
    ['指标名称', '目标值', '评估周期', '责任部门'],
    [
        ['年度培训覆盖率', '≥ 100%', '年度', '人力资源部'],
        ['新员工岗位胜任周期', '≤ 90天', '月度', '人力资源部+用人部门'],
        ['关键岗位证书持有率', '≥ 95%', '半年度', '人力资源部'],
        ['培训满意度（平均分）', '≥ 4.2 / 5', '每期结束', '人力资源部'],
        ['培训后绩效提升率', '≥ 85%', '培训后90天', '用人部门'],
        ['内训课程开发数量', '≥ 12门/年', '年度', '人力资源部'],
    ],
    col_widths=[5, 3, 2.5, 3.5]
)

page_break()

# ========== 第二章 ==========
heading1('第二章  培训组织与职责')

heading2('第六条  组织架构')
body('公司培训管理采用"人力资源部统筹 + 直线经理负责 + 员工自主"的三级管理模式：')
add_table(
    ['层级', '责任主体', '核心职责'],
    [
        ['决策层', '总经理/管理层', '审批年度培训计划，审批重大培训预算，监督培训战略执行'],
        ['统筹层', '人力资源部', '制度建设、计划制定、资源整合、效果评估、档案管理'],
        ['执行层', '直线经理（部门负责人）', '培训需求提报、本部门培训实施、效果跟踪反馈'],
        ['主体层', '员工本人', '主动学习、参训考勤、学以致用、经验内化分享'],
    ],
    col_widths=[2, 4, 8]
)

heading2('第七条  人力资源部职责')
bullet('制定公司年度培训计划并组织实施')
bullet('建立并维护培训管理体系，编制/修订培训制度')
bullet('开发和管理内部培训课程及教材')
bullet('选拔、培养、激励内部培训讲师')
bullet('管理培训预算，规范培训费用报销流程')
bullet('对子公司培训工作进行指导、监督与考核')
bullet('建立员工培训档案，追踪培训效果')

heading2('第八条  直线经理（部门负责人）职责')
bullet('识别并提报本部门员工的培训需求')
bullet('参与内训课程开发和案例萃取')
bullet('对下属参加培训进行审批并跟踪训后应用')
bullet('配合进行在岗技能考核和任职资格评审')

heading2('第九条  员工权利与义务')
add_table(
    ['权利', '义务'],
    [
        ['参加公司组织的各类培训', '遵守培训纪律，完成培训任务'],
        ['享有培训费用报销（符合规定）', '达到培训后的服务期要求'],
        ['对培训内容提出改进建议', '将所学知识应用于工作实践'],
        ['申请外部培训（经审批）', '参与内部知识分享和经验传承'],
        ['查看个人培训档案记录', '配合培训效果评估和考核'],
    ],
    col_widths=[7, 7]
)

page_break()

# ========== 第三章 ==========
heading1('第三章  培训体系架构')

heading2('第十条  培训体系总览')
body('公司培训体系覆盖员工职业全周期，分为四大培训类别和五个专业岗位序列：')

heading3('（一）四大培训类别')
add_table(
    ['培训类别', '定位', '触发机制', '责任部门'],
    [
        ['新员工培训', '快速融入，胜任岗位', '入职即启动，系统化实施90天', '人力资源部主导'],
        ['在职技能提升', '持续精进，专业深化', '年度计划 + 实时需求', '人力资源部+业务部门'],
        ['管理层培训', '领导力与战略思维', '年度/半年度集中培训', '人力资源部主导'],
        ['外部培训取证', '资质获取，视野拓展', '按需申请，审批后执行', '人力资源部统筹'],
    ],
    col_widths=[3, 4.5, 4, 2.5]
)

heading3('（二）五类岗位序列及核心能力')
add_table(
    ['序列', '代表岗位', '核心培训内容', '关键证书'],
    [
        ['专业技术序列', '造价工程师、监理工程师、BIM工程师', '清单计价、定额套用、质量管控、BIM建模', '一级造价工程师、监理工程师、BIM建模师'],
        ['项目管理序列', '全过程咨询顾问、项目经理', '项目管理、合同管理、设计管理、风险管理', '一级建造师、PMP'],
        ['招标代理序列', '招标专员、招标项目经理', '招投标法规、流程合规、文件编制、争议处理', '招标师'],
        ['职能支持序列', '财务、人事、行政、法务', '专业技能、上市合规、职业素养、数据分析', '对应专业职称'],
        ['高层管理序列', '总经理、副总经理、总监', '战略规划、行业政策、领导力、资本运营', '高级职称'],
    ],
    col_widths=[2.5, 4, 5, 2.5]
)

heading2('第十一条  培训方式')
add_table(
    ['培训方式', '适用场景', '特点'],
    [
        ['内部集中培训', '新员工通识、专项技能、管理层培训', '系统性强，成本低，便于统一管理'],
        ['岗位导师制', '新员工在岗辅导，专业序列传承', '个性化，针对性强，实战指导'],
        ['在线学习（E-Learning）', '法规学习、基础知识、碎片化学习', '灵活便捷，可重复，随时随地'],
        ['项目实践', '全过程咨询、BIM等实操技能', '以真实项目为载体，学以致用'],
        ['外派进修', '高管培训、资质取证、高端论坛', '视野拓展，资源整合，激励人才'],
        ['标杆交流', '管理层，行业考察', '对标学习，拓展思路'],
        ['读书会/分享会', '全员，管理层', '知识共享，营造学习氛围'],
    ],
    col_widths=[3, 4.5, 6.5]
)

heading2('第十二条  年度培训日历')
body('公司建立年度培训日历制度，每月固定培训项目如下：')
add_table(
    ['月份', '固定培训项目', '对象'],
    [
        ['1月', '年度培训计划发布会 + 开门红培训', '全员'],
        ['2月', '开工安全质量培训', '全员'],
        ['3月', '新员工90天培养启动', '新入职员工'],
        ['4月', 'BIM技术专项培训', '技术序列'],
        ['5月', '全过程咨询能力提升', '项目序列'],
        ['6月', '招投标法规更新培训（含年度继续教育）', '招标序列+全员'],
        ['7月', '半年培训效果复盘 + 经验分享会', '管理层'],
        ['8月', '中层管理能力提升', '部门负责人'],
        ['9月', '新员工90天培养结业答辩', '新员工'],
        ['10月', '行业政策解读专题', '管理层'],
        ['11月', '年度培训满意度调查 + 次年需求调研', '全员'],
        ['12月', '年度总结表彰 + 次年培训计划发布', '全员'],
    ],
    col_widths=[1.5, 8, 4.5]
)

page_break()

# ========== 第四章 ==========
heading1('第四章  新员工培训（90天培养计划）')

heading2('第十三条  培养目标')
body('通过90天系统化培养，使新员工：')
bullet('认同公司文化，理解公司战略和价值观')
bullet('掌握岗位必备的基础知识和基本技能')
bullet('能够独立完成本岗位基础性工作')
bullet('融入团队，建立职业人际网络')

heading2('第十四条  培养阶段总览')
add_table(
    ['阶段', '周期', '核心任务', '培养方式', '考核方式', '责任人'],
    [
        ['通识培训', '第1周', '文化融入、规章制度、基础业务认知', '集中授课+参观学习', '笔试（≥60分）', '人力资源部'],
        ['轮岗认知', '第2-4周', '主要业务板块快速认知', '各部门轮转+负责人座谈', '轮岗心得报告', '人力资源部'],
        ['岗位实训', '第5-12周', '本岗位核心技能学习', '导师制+项目实战', '实操考核+周报', '直线经理'],
        ['定岗考核', '第13周', '综合能力评定', '独立承担基础任务', '答辩评审', '人力资源部+部门'],
    ],
    col_widths=[2, 1.8, 3.5, 2.8, 2.5, 2.4]
)

heading2('第十五条  各阶段详细安排')

heading3('第一阶段：通识培训（第1周）')
add_table(
    ['课程模块', '课时', '内容要点', '授课人'],
    [
        ['公司介绍与文化', '4', '集团概况、发展历程、核心价值观、企业精神', '人力资源部负责人'],
        ['战略与业务全景', '4', '主营业务（造价/监理/招标/全过程）、商业模式、核心竞争力', '市场部/业务负责人'],
        ['组织架构与制度', '2', '组织架构图、汇报线、核心制度（财务/人事/合规）', '人力资源部'],
        ['行业基础知识', '4', '工程造价行业概述、全过程咨询、BIM技术简介', '业务部门负责人'],
        ['职场礼仪与素养', '2', '职业形象、沟通技巧、时间管理', '人力资源部'],
        ['IT系统操作', '2', 'OA系统、项目管理系统、造价软件基础', 'IT支持'],
        ['消防与安全教育', '2', '职场安全、信息安全、保密合规', '综合管理部'],
    ],
    col_widths=[3.5, 1.5, 6, 3]
)

heading3('第二阶段：轮岗认知（第2-4周）')
body('新员工在导师带领下，对公司核心业务板块进行认知轮转：')
add_table(
    ['轮转部门', '轮转时长', '学习目标', '产出成果'],
    [
        ['造价咨询部', '1周', '了解造价业务流程、定额体系、成果文件', '业务流程图笔记'],
        ['招标代理部', '3天', '了解招标流程、法规要求、文件类型', '招标流程图笔记'],
        ['工程监理部', '3天', '了解监理职责、现场管理、质量控制', '监理要点理解笔记'],
        ['全过程项目组', '3天', '了解全过程咨询各阶段、协同方式', '全流程认知报告'],
        ['BIM中心', '2天', '了解BIM应用场景、建模基础', 'BIM应用场景理解'],
    ],
    col_widths=[3.5, 2, 4.5, 4]
)

heading3('第三阶段：岗位实训（第5-12周）')
body('本阶段采用"导师制 + 项目实战"的双轨培养模式：')
bullet('每位新员工配备1名导师（原则上为同组高级工程师或项目经理）')
bullet('导师为新员工制定个人实训计划（IDP， Individual Development Plan）')
bullet('每周提交《周报》至人力资源部和导师，汇报学习进展')
bullet('每月接受1次阶段性考核，考核结果纳入转正评价')

heading3('第四阶段：定岗考核（第13周）')
body('新员工在第90天前完成答辩评审：')
bullet('提交《新员工培训总结报告》（不少于2000字）')
bullet('答辩委员会：人力资源部负责人 + 直线经理 + 导师')
bullet('答辩内容：岗位技能展示 + 典型案例分析 + 问题回答')
bullet('评分维度：业务理解（30%）、技能掌握（40%）、态度素养（30%）')
bullet('考核结果：优秀/良好/合格/不合格，决定定岗及薪资调整')

heading2('第十六条  导师管理')
add_table(
    ['导师职责', '导师激励', '考核要求'],
    [
        ['制定新员工IDP计划', '享受导师补贴：200元/人/月', '每月提交辅导记录不少于2次'],
        ['业务指导和问题答疑', '年度优秀导师评选，颁发荣誉证书', '辅导期离职率≤30%'],
        ['定期沟通反馈', '累计辅导满3人优先晋升加分', '新员工转正通过率≥90%'],
    ],
    col_widths=[4.5, 4.5, 5]
)

heading2('第十七条  淘汰机制')
body('新员工在90天内出现以下情形，公司有权依法解除劳动合同：')
bullet('试用期考核不合格，且经补考或延长试用期后仍不合格')
bullet('严重违反公司规章制度或保密纪律')
bullet('发现其入职时提供的材料存在虚假陈述')
bullet('旷工连续3天或累计5天以上')

page_break()

# ========== 第五章 ==========
heading1('第五章  在职技能提升培训')

heading2('第十八条  年度培训计划制定流程')
body('每年12月启动次年培训需求调研，计划于次年1月15日前发布：')
numbered('各部门提交本部门年度培训需求表（HR-002）', level=1, num='①')
numbered('人力资源部汇总分析，结合公司战略确定培训重点项目', level=1, num='②')
numbered('人力资源部编制《年度培训计划草案》', level=1, num='③')
numbered('总经理办公会审批', level=1, num='④')
numbered('发布实施（人力资源部发文，各部门执行）', level=1, num='⑤')

heading2('第十九条  重点培训项目')
add_table(
    ['培训项目', '培训对象', '课时', '频次', '目标'],
    [
        ['清单计价与定额应用提升', '造价工程师', '16课时', '每年2次', '覆盖全部造价人员，持证率100%'],
        ['全过程工程咨询实战', '项目序列全员', '24课时', '每年1次', '核心人员全部轮训'],
        ['BIM建模与应用专项', '技术序列全员', '32课时', '每年1次', '持BIM证书率≥80%'],
        ['招投标法规与合规', '招标序列+相关人员', '8课时', '每年1次', '全员覆盖，合规零违规'],
        ['安全生产与质量管理', '监理序列+现场人员', '12课时', '每季度1次', '安全事故为零'],
        ['上市合规与信息披露', '职能序列+管理层', '8课时', '每年1次', '相关人员全部培训'],
        ['职业素养与沟通技巧', '全员', '4课时', '每年1次', '提升客户满意度'],
    ],
    col_widths=[4, 3, 2, 2, 3]
)

heading2('第二十条  内部讲师制度')
body('公司建立三级内部讲师体系：')
add_table(
    ['级别', '申报条件', '课酬标准', '年度任务', '晋升条件'],
    [
        ['初级讲师', '主管级及以上，岗位经验丰富', '200元/课时', '≥20课时/年', '累计30课时+内部评审'],
        ['中级讲师', '初级讲师满1年，授课评价≥4.0', '400元/课时', '≥30课时/年', '累计80课时+评审'],
        ['高级讲师', '中级讲师满1年，课程体系化', '600元/课时', '≥40课时/年', '累计150课时+终审'],
    ],
    col_widths=[2.2, 4, 2.3, 2.5, 3]
)

body('内部讲师享有优先外训机会、年度优秀讲师评选等激励。课程开发质量纳入年度绩效考核加分项。')

heading2('第二十一条  在线学习（E-Learning）平台')
body('公司鼓励员工利用碎片时间进行自主学习：')
bullet('员工每年应完成不少于20课时的在线学习')
bullet('在线学习完成率纳入年度绩效考核（占培训模块5%分值）')
bullet('在线学习平台由人力资源部统一管理，记录自动同步至员工培训档案')

heading2('第二十二条  经验萃取与案例库')
body('公司建立内部知识沉淀机制：')
bullet('每年各部门须萃取不少于2个典型案例，提交至公司案例库')
bullet('案例库按专业序列分类管理：造价案例库/监理案例库/招标案例库/管理案例库')
bullet('优秀案例作者享有案例采纳奖金（标准：500元/案例）')
bullet('案例用于内训课程开发、新员工培训等场景')

page_break()

# ========== 第六章 ==========
heading1('第六章  管理层培训')

heading2('第二十三条  培训对象')
body('管理层培训面向以下人员：')
bullet('一级：总经理、副总经理、总监（公司决策层）')
bullet('二级：部门负责人、项目总监（中层管理）')
bullet('三级：后备管理人才（高潜员工）')

heading2('第二十四条  培训内容')
add_table(
    ['管理级别', '培训重点', '年度课时要求', '推荐方式'],
    [
        ['一级（决策层）', '企业战略、资本运营、上市合规、行业政策、领导力', '≥48课时/年', '外派高端培训/EMBA/标杆考察'],
        ['二级（中层）', '团队管理、项目管理、跨部门协作、绩效管理、沟通', '≥36课时/年', '内训+外训+读书会'],
        ['三级（后备）', '职业素养、角色转变、基础管理技能、业务全局观', '≥24课时/年', '内训+导师制+项目历练'],
    ],
    col_widths=[2.5, 5.5, 2.5, 3.5]
)

heading2('第二十五条  管理培训形式')
bullet('年度集中培训：公司每年组织1-2次管理层集中培训（2-3天/次）')
bullet('外派高端研修：选派优秀管理人员参加外部管理课程（MBA、总裁班、行业论坛等）')
bullet('标杆企业考察：每年组织1-2次赴行业标杆企业或优秀客户单位参访交流')
bullet('管理层读书会：每月至少读1本管理类书籍，每季度举办1次读书分享会')
bullet('行动学习项目：每年确定1个管理改善课题，以小组形式攻关解决')

heading2('第二十六条  外训审批流程')
numbered('参训人员提交《外部培训申请表》（HR-003）及课程资料', level=1, num='①')
numbered('直线经理初审，评估业务关联度和必要性', level=1, num='②')
numbered('人力资源部复审，核查费用预算', level=1, num='③')
numbered('总经理审批（单次费用≥1万元须总经理办公会审批）', level=1, num='④')
numbered('培训结束后5个工作日内提交《培训报告》及发票报销', level=1, num='⑤')

page_break()

# ========== 第七章 ==========
heading1('第七章  外部培训与证书管理')

heading2('第二十七条  证书分类管理')
body('公司所需执业资格证书按重要程度分为A、B两类：')
add_table(
    ['类别', '证书范围', '管理要求'],
    [
        ['A类（公司战略必需）', '一级造价工程师、一级建造师、监理工程师', '优先安排培训，100%持证目标，强制持有'],
        ['B类（业务发展鼓励）', '招标师、二级造价工程师、BIM高级证书', '按需培训，鼓励持有，给予补贴'],
    ],
    col_widths=[3, 5.5, 5.5]
)

heading2('第二十八条  培训费用报销标准')
add_table(
    ['证书类型', '公司报销上限', '服务期要求', '违约处理'],
    [
        ['一级建造师', '8000元/人', '取得后服务满3年', '未满服务期离职，按比例退还培训费'],
        ['一级造价工程师', '6000元/人', '取得后服务满3年', '同上'],
        ['监理工程师', '4000元/人', '取得后服务满2年', '同上'],
        ['招标师', '3000元/人', '取得后服务满2年', '同上'],
        ['BIM高级建模师', '3000元/人', '取得后服务满2年', '同上'],
        ['其他执业资格', '2000元/人', '取得后服务满1年', '同上'],
    ],
    col_widths=[3.5, 3, 3, 4.5]
)

heading2('第二十九条  证书技能补贴')
body('员工取得公司所需执业资格证书后，除报销培训费用外，享受月度技能补贴：')
add_table(
    ['证书', '一次性奖励', '月度技能补贴', '备注'],
    [
        ['一级建造师', '5000元', '500元/月', '同时担任项目经理另议'],
        ['一级造价工程师', '4000元', '400元/月', '专业带头人另议'],
        ['监理工程师', '3000元', '300元/月', '担任总监另议'],
        ['招标师', '2000元', '200元/月', '—'],
        ['BIM高级建模师', '2000元', '200元/月', '—'],
    ],
    col_widths=[3.5, 2.5, 2.5, 5.5]
)

heading2('第三十条  继续教育管理')
bullet('执业资格证书的继续教育费用由公司全额报销')
bullet('继续教育的形式包括：线上课程、面授培训、行业会议学时等')
bullet('员工应在证书有效期满前完成继续教育并换证，确保证书持续有效')
bullet('人力资源部负责建立证书有效期台账，提前6个月提醒相关人员参加继续教育')

page_break()

# ========== 第八章 ==========
heading1('第八章  培训效果评估体系')

heading2('第三十一条  柯氏四级评估模型')
body('公司采用国际通用的柯氏四级评估模型（Kirkpatrick Model）进行培训效果评估：')

add_table(
    ['评估层级', '评估内容', '评估方式', '评估时点', '数据来源'],
    [
        ['第一级\n反应层', '学员满意度\n培训体验', '满意度调查问卷\n（5分制）', '培训结束时', '学员自评'],
        ['第二级\n学习层', '知识技能\n掌握程度', '笔试+实操考核\n（≥60分及格）', '培训结束\n+30天', '人力资源部'],
        ['第三级\n行为层', '训后行为\n应用程度', '直线经理评价\n+360度反馈', '培训后\n90天', '直线经理\n+同事/下属'],
        ['第四级\n结果层', '对组织业绩\n的贡献', '绩效数据对比\n培训ROI分析', '培训后\n6-12个月', '绩效系统\n财务数据'],
    ],
    col_widths=[2.5, 3, 4, 2.5, 2]
)

heading2('第三十二条  评估结果应用')
bullet('评估结果记入员工培训档案，作为晋升、评优的重要参考')
bullet('培训满意度连续2期低于3.5分的课程，暂停开设并优化改进')
bullet('培训后行为转化率达标的员工，优先安排下一阶段培训')
bullet('培训投入产出比（ROI）作为年度培训预算调整依据')

heading2('第三十三条  培训满意度调查表')
body('每次培训结束时发放《培训满意度调查问卷》（HR-005），评估维度包括：')
add_table(
    ['评估维度', '权重', '优秀标准（5分）'],
    [
        ['课程内容实用性', '25%', '4.5分以上'],
        ['讲师授课质量', '20%', '4.3分以上'],
        ['培训组织管理', '15%', '4.0分以上'],
        ['学习收获感', '20%', '4.2分以上'],
        ['建议与意见（开放）', '20%', '定性收集'],
    ],
    col_widths=[4, 2, 8]
)

page_break()

# ========== 第九章 ==========
heading1('第九章  培训纪律与考核')

heading2('第三十四条  培训出勤管理')
bullet('员工参加培训须提前5分钟到达培训场地，因故不能参加者须提前填写《培训请假单》（HR-004）')
bullet('培训期间手机须调至静音或震动模式，不得在培训期间处理与培训无关的事务')
bullet('培训期间迟到10分钟以上者，本次培训不记学时；旷课者视为培训不合格')

heading2('第三十五条  请假审批权限')
add_table(
    ['请假类型', '审批人', '备注'],
    [
        ['培训时间≤4课时', '直线经理', '须提前1天申请'],
        ['培训时间4-16课时', '直线经理+人力资源部', '须提前3天申请'],
        ['培训时间>16课时（外训）', '直线经理+人力资源部+分管领导', '须提前1周申请'],
    ],
    col_widths=[4, 5, 5]
)

heading2('第三十六条  培训违规处理')
add_table(
    ['违规情形', '处理方式'],
    [
        ['无故旷课', '扣罚当月绩效50元/课时，记入培训档案，书面警告'],
        ['培训考核不合格（不补考）', '扣发当月全勤奖，不得参加当年评优'],
        ['伪造培训记录/学时', '解除劳动合同，并追究法律责任'],
        ['泄露公司培训资料（涉密）', '解除劳动合同，并追究法律责任'],
    ],
    col_widths=[5, 9]
)

heading2('第三十七条  培训考核结果处理')
add_table(
    ['考核结果', '处理方式'],
    [
        ['优秀（≥90分）', '通报表扬，记入培训档案，优先安排下一阶段培训'],
        ['良好（75-89分）', '肯定记录，记入培训档案'],
        ['合格（60-74分）', '记录在案，要求补学相关章节'],
        ['不合格（<60分）', '补考一次，补考仍不合格者换岗或降薪处理'],
    ],
    col_widths=[3, 11]
)

page_break()

# ========== 第十章 ==========
heading1('第十章  培训档案管理')

heading2('第三十八条  档案内容')
body('人力资源部为每位员工建立《员工培训档案》（HR-007），档案内容包含：')
bullet('个人基本信息：姓名、部门、岗位、入职日期、岗位序列')
bullet('培训记录：课程名称、培训形式、培训日期、课时、考核成绩、出勤情况')
bullet('证书记录：证书名称、取得日期、有效期、注册状态、继续教育完成情况')
bullet('培训费用：本人发生的内部培训成本、外部培训报销金额合计')
bullet('内训贡献：担任内训讲师情况、开发课程数量、案例贡献')

heading2('第三十九条  档案管理规范')
bullet('纸质档案与电子档案并行管理，确保信息安全')
bullet('档案保存期限：员工离职后不少于3年')
bullet('员工有权查阅本人培训档案，每年可申请1次纸质复印件')
bullet('员工离职时，人力资源部出具《员工培训证明》，注明培训经历和持证情况')

heading2('第四十条  培训数据月度报送')
body('人力资源部每月向总经理报送《培训工作月报》，内容包括：')
bullet('本月开展培训场次、覆盖人数、总课时')
bullet('本月培训费用支出明细')
bullet('本月新员工培养进度跟踪')
bullet('本月持证情况更新')
bullet('下月培训计划')

page_break()

# ========== 第十一章 ==========
heading1('第十一章  年度培训实施计划表')

heading2('第四十一条  2025年度重点培训计划')
body('依据公司2025年战略目标，本年度重点培训计划如下：')

add_table(
    ['序号', '培训项目', '培训对象', '计划月份', '计划课时', '负责人', '预算（元）', '预期成果'],
    [
        ['1', '新员工90天培养（全年滚动）', '全体新入职员工', '每月启动', '120课时/人', '人力资源部', '按人头分摊', '100%按期转正'],
        ['2', '清单计价与定额应用提升', '造价工程师（全员）', '3月、9月', '16课时', '造价部负责人', '20,000', '持证率100%'],
        ['3', '全过程工程咨询实战训练营', '项目序列全员', '5月', '24课时', '项目管理部', '30,000', '核心人员轮训完成'],
        ['4', 'BIM建模专项培训（初级+中级）', '技术序列', '4月、8月', '32课时', 'BIM中心', '25,000', 'BIM证书获取率80%'],
        ['5', '招投标法规与合规培训', '招标序列+相关', '6月', '8课时', '招标部', '8,000', '合规零违规'],
        ['6', '安全生产月专项培训', '监理序列+现场', '6月', '12课时', '安全质量部', '10,000', '安全事故为零'],
        ['7', '中层管理能力提升培训', '部门负责人', '8月', '16课时', '人力资源部', '40,000', '管理能力评估提升'],
        ['8', '新员工90天结业答辩', '本期新员工', '6月、9月、12月', '答辩', '人力资源部', '0', '人才识别'],
        ['9', '年度培训满意度调研', '全员', '11月', '问卷', '人力资源部', '2,000', '满意度≥4.2分'],
        ['10', '年度培训总结表彰', '全员', '12月', '半天', '人力资源部', '15,000', '激励先进'],
    ],
    col_widths=[0.8, 4, 3, 1.8, 1.5, 2.5, 1.8, 3.6]
)

heading2('第四十二条  培训预算')
body('年度培训预算按以下标准提取和使用：')
add_table(
    ['预算项目', '计算标准', '使用范围', '审批权限'],
    [
        ['工资总额计提', '工资总额的1.5%', '全部培训支出', '年度预算内，总经理审批'],
        ['新员工培训专项', '按实际发生列支', '入职培训、导师补贴', '纳入年度培训预算'],
        ['外训支出', '按实际报销', '外部培训费、差旅费', '单次≥1万元，总经理办公会'],
        ['讲师课酬', '按标准发放', '内部讲师课酬', '人力资源部审批'],
        ['证书奖励', '按标准发放', '一次性奖励、月度补贴', '人力资源部审批，财务执行'],
    ],
    col_widths=[2.5, 2.5, 4, 5]
)

page_break()

# ========== 第十二章 ==========
heading1('第十二章  表单清单与填写说明')

heading2('第四十三条  表单目录')
add_table(
    ['表单编号', '表单名称', '使用场景', '填制节点', '责任主体'],
    [
        ['HR-001', '新员工入职培训记录表', '新员工入职培训全过程记录', '入职第1周', '人力资源部'],
        ['HR-002', '员工年度培训需求申请表', '各部门提报次年培训需求', '每年12月', '各部门负责人'],
        ['HR-003', '外部培训申请表', '员工申请参加外训', '培训前1周', '申请人+各级审批'],
        ['HR-004', '培训请假单', '员工因故无法参加培训', '培训前', '申请人+直线经理'],
        ['HR-005', '培训满意度调查问卷', '每期培训结束时收集', '培训结束时', '人力资源部'],
        ['HR-006', '培训考核评分表', '培训结束时进行考核', '培训结束时', '讲师+人力资源部'],
        ['HR-007', '员工培训档案表', '一人一档，动态更新', '随时更新', '人力资源部'],
        ['HR-008', '内部讲师课酬申请表', '内部讲师申请课酬', '每月末', '讲师+人力资源部'],
        ['HR-009', '证书补贴申请表', '员工取得证书后申领补贴', '取证后1个月内', '申请人+人力资源部'],
        ['HR-010', '导师辅导记录表', '导师记录辅导期沟通', '每月至少2次', '导师'],
        ['HR-011', '培训效果评估表（柯氏四级）', '培训后30/90天效果追踪', '培训后30/90天', '人力资源部+直线经理'],
        ['HR-012', '新员工周报', '新员工每周汇报学习进展', '每周五', '新员工+导师'],
        ['HR-013', '新员工IDP计划表', '个人发展计划', '入职第5周前', '导师+新员工+人力资源部'],
        ['HR-014', '新员工培训总结报告', '90天培训成果汇报', '入职第85天前', '新员工'],
        ['HR-015', '年度培训计划表', '各部门提报年度培训计划', '每年12月', '人力资源部'],
        ['HR-016', '培训工作月报', '每月培训工作总结', '每月5日前', '人力资源部'],
        ['HR-017', '培训签到表', '每次培训记录出席情况', '培训进行中', '人力资源部/讲师'],
        ['HR-018', '培训费用报销单', '培训费用报销申请', '培训结束后5个工作日', '申请人+财务'],
    ],
    col_widths=[2, 4.5, 4, 2.5, 2.5]
)

heading2('第四十四条  表单电子化管理')
body('上述表单除HR-001、HR-007须保留纸质原件外，其余表单均采用电子化管理：')
bullet('电子表单统一存放于OA系统"人力资源-培训管理"模块')
bullet('表单编号由系统自动生成，确保唯一性')
bullet('所有审批流程在线完成，支持电子签名')
bullet('人力资源部每季度对表单填写质量进行检查')

page_break()

# ========== 第十三章 ==========
heading1('第十三章  附则')

heading2('第四十五条  解释权')
body('本制度由人力资源部负责解释。在执行过程中如有疑问或争议，由人力资源部提请总经理办公会裁决。')

heading2('第四十六条  修订机制')
body('本制度每年修订一次。有下列情形之一的，可即时修订：')
bullet('公司组织架构发生重大调整')
bullet('国家或地方相关法律法规修订')
bullet('公司战略方向发生重大变化')
bullet('执行中发现制度存在重大疏漏')

heading2('第四十七条  生效')
body('本制度自2025年1月1日起施行。原《培训管理制度》（版本号V0.X）同步废止。')

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('中诚智信工程咨询集团股份有限公司')
set_font(run, '宋体', 11, bold=True, color=BRAND_DARK)
set_para_spacing(p, before=6, after=4)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run2 = p2.add_run('制定部门：人力资源部')
set_font(run2, '宋体', 10, color=TEXT_GREY)
set_para_spacing(p2, before=0, after=4)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run3 = p3.add_run('2025年1月  第一版发布')
set_font(run3, '宋体', 10, color=TEXT_GREY)
set_para_spacing(p3, before=0, after=4)

doc.add_paragraph()
divider()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('— 附录：本制度配套表单编号 HR-001 至 HR-018，详见第五章第四十三条 —')
set_font(run4, '宋体', 9, italic=True, color=TEXT_GREY)
set_para_spacing(p4, before=6, after=6)

# ========== 保存 ==========
output_path = '/Users/huahaha/Downloads/中诚咨询培训管理制度_V1.0_2025.docx'
doc.save(output_path)
print(f'✅ 文档已生成: {output_path}')
