#!/usr/bin/env python3
"""合并两个培训制度文档"""
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

# 加载两个文档
doc1 = Document('/Users/huahaha/Downloads/中诚咨询全套培训制度_完整版.docx')
doc2 = Document('/Users/huahaha/Downloads/中诚咨询全套培训制度_完整版_Part2.docx')

# 创建一个新文档
merged = Document()

# 复制 doc1 的所有内容
for element in doc1.element.body:
    merged.element.body.append(deepcopy(element))

# 复制 doc2 的所有内容（从第一个元素开始）
for element in doc2.element.body:
    merged.element.body.append(deepcopy(element))

# 保存
output = '/Users/huahaha/Downloads/中诚咨询全套培训管理制度_完整版_最终.docx'
merged.save(output)
print(f'✅ 合并完成: {output}')

import os
size = os.path.getsize(output)
print(f'   文件大小: {size/1024:.1f} KB')
